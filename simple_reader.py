import docx

def read_all_text(filename):
    """Reads and prints all text from a .docx file."""
    print(f"Attempting to read text from: {filename}")
    try:
        doc = docx.Document(filename)

        print("\n--- Reading Paragraphs ---\n")
        full_text = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"Para {i}: {para.text}")
                full_text.append(para.text)

        if not full_text:
            print("No text found in paragraphs.")

        print("\n--- Reading Tables ---\n")
        table_text_found = False
        if doc.tables:
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            print(f"Table {t_idx}, Row {r_idx}, Cell {c_idx}: {cell.text}")
                            table_text_found = True

        if not table_text_found:
            print("No text found in tables.")

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    read_all_text('Contenus site Hawaii.docx')