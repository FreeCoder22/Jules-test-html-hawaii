import docx
import json
import re

def get_full_text(doc):
    """Extracts all text from the document, including from tables."""
    text_lines = []
    for para in doc.paragraphs:
        text_lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_lines.append(cell.text)
    return "\n".join(text_lines)

def convert_docx_to_json(docx_path, json_path):
    print(f"Opening DOCX file: {docx_path}")
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        print(f"Error opening DOCX file: {e}")
        return

    full_text = get_full_text(doc)
    # Replace non-breaking spaces with regular spaces for easier regex matching
    full_text = full_text.replace('\xa0', ' ')

    data = []
    # Regex to split the text by page markers. This captures the filename.
    # It looks for 'PAGE' optionally followed by 'CIBLE', a colon, and then the HTML filename.
    pages_raw = re.split(r'PAGE\s*(?:CIBLE)?\s*:\s*(.*?\.html)', full_text, flags=re.IGNORECASE)

    # The first item in pages_raw is anything before the first match, which we can ignore.
    # The rest of the list is alternating [filename, content, filename, content, ...]
    page_data = pages_raw[1:]

    if not page_data:
        print("Warning: No page markers found. Please check the document's format.")
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump([], f)
        return

    # Iterate through the pairs of (filename, content)
    for i in range(0, len(page_data), 2):
        page_title = page_data[i].strip()
        page_content = page_data[i+1]

        print(f"Processing PAGE: {page_title}")

        current_page = {'page_cible': page_title, 'sections': []}

        # Regex to split the page's content by section markers. Captures the section title.
        sections_raw = re.split(r'SECTION\s*:\s*(.*?)\n', page_content, flags=re.IGNORECASE)

        # Ignore content before the first section and process pairs of (title, content)
        section_data = sections_raw[1:]

        for j in range(0, len(section_data), 2):
            section_title = section_data[j].strip()
            section_content_raw = section_data[j+1].strip()

            # Clean up the section content: split by newline, strip whitespace, and remove empty lines or closing braces
            cleaned_content = [
                line.strip() for line in section_content_raw.split('\n')
                if line.strip() and '}' not in line
            ]

            if not cleaned_content:
                continue

            print(f"  - Found SECTION: {section_title}")

            current_section = {'titre': section_title, 'contenu': cleaned_content}
            current_page['sections'].append(current_section)

        if current_page['sections']:
            data.append(current_page)

    print(f"Successfully extracted data for {len(data)} pages.")
    if not data:
        print("Warning: No data was ultimately extracted into the final JSON.")

    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

    print(f"Conversion to JSON complete. Output saved to {json_path}")

if __name__ == "__main__":
    convert_docx_to_json('Contenus site Hawaii.docx', 'contenu.json')